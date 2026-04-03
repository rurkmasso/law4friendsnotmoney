"""
Legal document drafter — generates DOCX contracts, letters, opinions in
Maltese and/or English, grounded in Malta law.
"""
import io
import anthropic
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sqlalchemy.ext.asyncio import AsyncSession
from rag.retriever import retrieve, build_context
from config import settings

_anthropic = anthropic.AsyncAnthropic(api_key=settings.anthropic_api_key)

TEMPLATES = {
    "sale_agreement": "Sale of Immovable Property Agreement",
    "employment_contract": "Contract of Employment",
    "service_agreement": "Service Agreement",
    "nda": "Non-Disclosure Agreement",
    "company_resolution": "Board Resolution",
    "power_of_attorney": "Power of Attorney",
    "lease_agreement": "Lease Agreement",
    "loan_agreement": "Loan Agreement",
    "partnership_deed": "Partnership Deed",
    "will": "Last Will and Testament",
    "demand_letter": "Formal Demand Letter",
    "legal_opinion": "Legal Opinion",
}

DRAFT_SYSTEM_EN = """You are a Maltese legal drafter. Draft professional legal documents governed by Maltese law.
- Use correct Maltese legal terminology
- Cite the relevant chapters of the Laws of Malta
- Include standard protective clauses under Maltese law
- Format with clear numbered clauses
- Output only the document body (no meta-commentary)"""

DRAFT_SYSTEM_MT = """Inti abbozzatur legali Malti. Abbozza dokumenti legali professjonali regolati mil-liġi Maltija.
- Uża t-terminoloġija legali Maltija korretta
- Ċita l-kapitoli rilevanti tal-Liġijiet ta' Malta
- Inkludi klawżoli protettivi standard skont il-liġi Maltija
- Formatta b'klawżoli numerati ċari
- Ipproduċi biss il-korp tad-dokument"""


async def draft_document(
    doc_type: str,
    instructions: str,
    language: str,
    db: AsyncSession,
) -> bytes:
    """Returns DOCX file bytes."""
    template_title = TEMPLATES.get(doc_type, doc_type.replace("_", " ").title())

    # Retrieve relevant laws to ground the draft
    chunks = await retrieve(f"{template_title} Malta law", db)
    context = build_context(chunks)

    system = DRAFT_SYSTEM_MT if language == "mt" else DRAFT_SYSTEM_EN
    lang_label = "Malti" if language == "mt" else "English"

    prompt = f"""Relevant Malta law context:
{context}

---
Draft a {template_title} in {lang_label}.
Instructions: {instructions}

Output the full document text with numbered clauses."""

    response = await _anthropic.messages.create(
        model=settings.model_complex,  # always Sonnet for drafting
        max_tokens=4096,
        system=system,
        messages=[{"role": "user", "content": prompt}],
    )

    document_text = response.content[0].text
    return _build_docx(template_title, document_text, language)


def _build_docx(title: str, body: str, language: str) -> bytes:
    doc = DocxDocument()

    # Styling
    style = doc.styles["Normal"]
    style.font.name = "Georgia"
    style.font.size = Pt(11)

    # Header
    header = doc.add_heading(title.upper(), level=1)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    lang_note = "Abbozzat skont il-Liġi Maltija" if language == "mt" else "Drafted under the Laws of Malta"
    sub = doc.add_paragraph(lang_note)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(9)
    sub.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # Body — split on numbered clauses
    for line in body.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        # Bold clause headings like "1." or "CLAUSE 1"
        import re
        if re.match(r"^(\d+\.|CLAUSE \d+|Article \d+)", line, re.I):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
        else:
            doc.add_paragraph(line)

    # Footer
    doc.add_paragraph()
    footer_text = (
        "Dan id-dokument ġie abbozzat bl-għajnuna ta' LexMalta AI. "
        "Mhux parir legali. Ikkonsulta avukat." if language == "mt" else
        "This document was drafted with LexMalta AI assistance. "
        "Not legal advice. Consult a warranted advocate."
    )
    footer = doc.add_paragraph(footer_text)
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
